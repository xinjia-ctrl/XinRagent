from pathlib import Path
from typing import Any

from app.core.exceptions import RagentException
from app.ingestion.context import ParsedDocument
from app.ingestion.parser.base import DocumentParser


class PdfParser(DocumentParser):
    supported_types = {"pdf"}

    async def parse(self, path: Path) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise _missing_dependency("pypdf", "PDF") from exc

        reader = PdfReader(str(path))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"## 第 {index} 页\n{text.strip()}")
        return ParsedDocument(
            text="\n\n".join(pages),
            metadata={"parser": "pdf", "pageCount": len(reader.pages)},
        )


class DocxParser(DocumentParser):
    supported_types = {"docx"}

    async def parse(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise _missing_dependency("python-docx", "DOCX") from exc

        document = Document(str(path))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return ParsedDocument(
            text="\n".join(blocks),
            metadata={"parser": "docx", "paragraphCount": len(document.paragraphs), "tableCount": len(document.tables)},
        )


class XlsxParser(DocumentParser):
    supported_types = {"xlsx", "xlsm"}

    async def parse(self, path: Path) -> ParsedDocument:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise _missing_dependency("openpyxl", "XLSX") from exc

        workbook = load_workbook(str(path), read_only=True, data_only=True)
        try:
            sections = []
            for sheet in workbook.worksheets:
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    values = [_cell_to_text(value) for value in row]
                    if any(values):
                        rows.append(" | ".join(values).strip())
                if rows:
                    sections.append(f"## {sheet.title}\n" + "\n".join(rows))
            return ParsedDocument(
                text="\n\n".join(sections),
                metadata={"parser": "xlsx", "sheetCount": len(workbook.worksheets)},
            )
        finally:
            workbook.close()


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _missing_dependency(package_name: str, file_label: str) -> RagentException:
    return RagentException(
        message=f"解析 {file_label} 需要安装依赖 {package_name}",
        code="INGESTION_PARSER_DEPENDENCY_MISSING",
        status_code=500,
    )
